import pandas as pd
from langchain.tools import Tool
import os
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os
import base64
from datetime import datetime
from sendgrid import SendGridAPIClient
from sendgrid.helpers.mail import (
    Mail, Email, To, Content, 
    Attachment, FileContent, FileName, FileType, Disposition
)
from dotenv import load_dotenv
from langchain.tools import StructuredTool

load_dotenv(override=True)

def get_greeting():
    """Returns a greeting based on the time of day."""
    hour = datetime.datetime.now().hour
    if hour < 12:
        return "Good Morning"
    elif hour < 17:
        return "Good Afternoon"
    else:
        return "Good Evening"

def set_cell_border(cell, **kwargs):
    """
    Set cell border properties.
    Usage: set_cell_border(cell, top={"sz": 12, "val": "single", "color": "000000"})
    """
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        if edge in kwargs:
            edge_data = kwargs[edge]
            element = OxmlElement(f'w:{edge}')
            for key in ["sz", "val", "color", "space"]:
                if key in edge_data:
                    element.set(qn(f'w:{key}'), str(edge_data[key]))
            tcBorders.append(element)
    tcPr.append(tcBorders)

def read_invoice_data(filename: str) -> str:
    try:
        # Resolve the full path
        filepath = os.path.join(os.getcwd(), filename)
        
        # Read the Excel file
        df = pd.read_excel(filepath)

        if df.empty:
            return "The Excel file is empty."

        # Convert the DataFrame into a nicely formatted table-like string
        lines = []
        for _, row in df.iterrows():
            date = row.get("Date", "")
            status = row.get("Status", "")
            remarks = row.get("Remarks", "")
            lines.append(f"{date} | {status} | {remarks}")

        return "Timesheet Records:\n" + "\n".join(lines)

    except Exception as e:
        return f"Error reading Excel timesheet: {str(e)}"

def create_invoice_document(filename: str, data: dict) -> str:
    """    Generates a .docx invoice from a dictionary of data.    """
    try:

        if not filename or not data:
            return "Error: Both 'filename' and 'data' are required."

        if not filename.endswith(".docx"):
            filename += ".docx"

        # Use an absolute path or a designated output directory
        output_dir = os.path.join(os.getcwd())
        file_path = os.path.join(output_dir, filename)
        
        doc = Document()

        # --- Set Default Font Style ---
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Times New Roman'
        font.size = Pt(11)

        # --- Main Table (1 column, used for layout) ---
        table = doc.add_table(rows=1, cols=1)
        table.autofit = False
        table.allow_autofit = False
        table.columns[0].width = Inches(6.5)

        # --- Row 1: INVOICE Title ---
        cell = table.cell(0, 0)
        p = cell.paragraphs[0]
        run = p.add_run("INVOICE")
        run.font.name = "Arial Black"
        run.font.size = Pt(28)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_border(cell, bottom={"val": "nil"}, left={"sz": 6, "val": "single"}, right={"sz": 6, "val": "single"}, top={"sz": 6, "val": "single"})

        # --- Row 2: Name (left) and Date (right) ---
        row_cell = table.add_row().cells[0]
        inner_table = row_cell.add_table(rows=1, cols=2)
        inner_table.columns[0].width = Inches(4.0)
        inner_table.columns[1].width = Inches(2.5)
        
        left_cell = inner_table.cell(0, 0)
        left_cell.paragraphs[0].add_run(data["name"]).bold = True
        left_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT

        right_cell = inner_table.cell(0, 1)
        right_cell.paragraphs[0].add_run(f"{data['date']}")
        right_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
        for cell in inner_table._cells:
            set_cell_border(cell, top={"val": "nil"}, bottom={"val": "nil"}, left={"val": "nil"}, right={"val": "nil"})
        set_cell_border(row_cell, bottom={"val": "nil"}, left={"sz": 6, "val": "single"}, right={"sz": 6, "val": "single"}, top={"val": "nil"})

        # --- Row 3: Bill To Block ---
        cell = table.add_row().cells[0]
        p = cell.paragraphs[0]
        p.add_run("Bill To:\n").bold = True
        for line in data["bill_to"]:
            p.add_run(f"    {line}\n")
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        set_cell_border(cell, top={"val": "nil"}, bottom={"val": "nil"}, left={"sz": 6, "val": "single"}, right={"sz": 6, "val": "single"})

        # --- Row 4: Details Table Headers ---
        row_cell = table.add_row().cells[0]
        header_table = row_cell.add_table(rows=1, cols=2)
        header_table.columns[0].width = Inches(5.0)
        header_table.columns[1].width = Inches(1.5)
        desc_cell = header_table.cell(0, 0)
        amt_cell = header_table.cell(0, 1)
        header_table.alignment = WD_TABLE_ALIGNMENT.LEFT
        header_table.autofit = False

        desc_cell.text = "DESCRIPTION"
        amt_cell.text = "AMOUNT"

        for cell in [desc_cell, amt_cell]:
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.runs[0].bold = True
            shading_elm = OxmlElement('w:shd')
            shading_elm.set(qn('w:fill'), "ff99cc")
            cell._tc.get_or_add_tcPr().append(shading_elm)
        set_cell_border(row_cell, left={"sz": 6, "val": "single"}, right={"sz": 6, "val": "single"})

        # --- Row 5: Details Table Content ---
        row_cell = table.add_row().cells[0]
        details_table = row_cell.add_table(rows=0, cols=2)
        details_table.columns[0].width = Inches(5.0)
        details_table.columns[1].width = Inches(1.5)

        cells = details_table.add_row().cells
        cells[0].text = data["salary_description"]
        cells[1].text = ""

        # Add itemized details, handling both dict and string formats for robustness
        for item in data["details"]:
            cells = details_table.add_row().cells
            if isinstance(item, dict):
                # Handle the new, preferred dictionary format
                cells[0].text = item.get("description", "")
                cells[1].text = item.get("amount", "")
            elif isinstance(item, str) and ':' in item:
                # Handle the old string format for backward compatibility
                key, value = item.split(":", 1)
                cells[0].text = key.strip()
                cells[1].text = value.strip()
            else:
                # Handle any other unexpected format
                cells[0].text = str(item)
                cells[1].text = ""
            cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        set_cell_border(row_cell, top={"val": "nil"}, bottom={"sz": 6, "val": "single"}, left={"sz": 6, "val": "single"}, right={"sz": 6, "val": "single"})

        # --- Row 6: TOTAL Row ---
        row_cell = table.add_row().cells[0]
        total_table = row_cell.add_table(rows=1, cols=2)
        total_table.columns[0].width = Inches(5.0)
        total_table.columns[1].width = Inches(1.5)
        left = total_table.cell(0, 0)
        right = total_table.cell(0, 1)

        left.paragraphs[0].add_run("TOTAL").bold = True
        left.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        right.paragraphs[0].add_run(data["total"]).bold = True
        right.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT

        shading_elm = OxmlElement('w:shd')
        shading_elm.set(qn('w:fill'), "ffcc99")
        right._tc.get_or_add_tcPr().append(shading_elm)
        right.alignment = WD_TABLE_ALIGNMENT.LEFT
        right.autofit = False

        set_cell_border(row_cell, top={"sz": 6, "val": "single"}, bottom={"val": "nil"}, left={"sz": 6, "val": "single"}, right={"sz": 6, "val": "single"})
        set_cell_border(right, left={"sz": 6, "val": "single"}, bottom={"sz": 6, "val": "single"})

        # --- Row 7: Amount in Words ---
        cell = table.add_row().cells[0]
        p = cell.paragraphs[0]
        p.add_run("Amount in Words: ").bold = True
        p.add_run(data["total_words"])
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        set_cell_border(cell, top={"val": "nil"}, bottom={"sz": 6, "val": "single"}, left={"sz": 6, "val": "single"}, right={"sz": 6, "val": "single"})

        doc.save(file_path)
        return f"Invoice successfully written to {file_path}"

    except Exception as e:
        return f"Failed to write Word document: {str(e)}"

def save_or_update_timesheet(filename: str, date: str, status: str, remarks: str) -> str:
    """
    Saves or updates a single entry in the Excel timesheet.
    If an entry for the given date already exists, it will be UPDATED.
    Otherwise, a new entry will be ADDED.
    The date must be in 'YYYY-MM-DD' format.
    """
    try:
        filepath = os.path.join(os.getcwd(), filename)
        
        # Define the columns to ensure consistency
        columns = ["Date", "Status", "Remarks"]

        # Check if the file exists and read it; otherwise, create an empty DataFrame
        if os.path.exists(filepath):
            df = pd.read_excel(filepath)
            # Ensure the Date column is treated as string to avoid formatting issues
            df['Date'] = df['Date'].astype(str).str.split(' ').str[0]
        else:
            df = pd.DataFrame(columns=columns)

        # Check if an entry for the given date already exists
        if date in df['Date'].values:
            # Update existing entry
            idx = df.index[df['Date'] == date][0]
            df.loc[idx, 'Status'] = status
            df.loc[idx, 'Remarks'] = remarks
            action = "updated"
        else:
            # Add new entry
            new_entry = pd.DataFrame([{"Date": date, "Status": status, "Remarks": remarks}])
            df = pd.concat([df, new_entry], ignore_index=True)
            action = "added"
            
        # Save the updated DataFrame back to Excel
        df.to_excel(filepath, index=False)
        
        return f"Success: The entry for {date} was {action} in {filename}."

    except Exception as e:
        return f"Error modifying Excel timesheet: {str(e)}"

def send_email_with_attachments(
    xlsx_filename: str, 
    docx_filename: str, 
    to_email: str = None
) -> str:
    """
    Sends an email with an XLSX and a DOCX file as attachments from the current directory.
    
    Args:
        xlsx_filename: The filename of the .xlsx file (e.g., 'timesheet_july.xlsx').
        docx_filename: The filename of the .docx file (e.g., 'invoice_july.docx').
        to_email: The recipient's email address. Defaults to the one in the .env file.
        
    Returns:
        A string indicating success or failure.
    """
    try:
        sg_api_key = os.getenv("SENDGRID_API_KEY")
        from_email_addr = os.getenv("FROM_EMAIL")
        to_email_addr = to_email or os.getenv("TO_EMAIL")

        if not all([sg_api_key, from_email_addr, to_email_addr]):
            raise ValueError("Missing required environment variables")

        greeting = get_greeting()
        message = Mail(
            from_email=Email(from_email_addr),
            to_emails=To(to_email_addr),
            subject=f"Timesheet and Invoice for Approval",
            plain_text_content=Content("text/plain", f"Hi,\n{greeting}.\n\nI've attached the timesheet and invoice for the month of July. Can review them and approve at your convenience.")
        )
        
        attachments = []
        file_info = {
            xlsx_filename: "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            docx_filename: "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        }

        for filename, mime_type in file_info.items():
            if os.path.exists(filename):
                with open(filename, "rb") as f:
                    data = f.read()
                encoded_file = base64.b64encode(data).decode()
                
                attached_file = Attachment(
                    FileContent(encoded_file),
                    FileName(filename),
                    FileType(mime_type),
                    Disposition("attachment")
                )
                attachments.append(attached_file)

        if not attachments:
            return "Failure: No valid files found in the current directory to attach."

        message.attachment = attachments

        sg = SendGridAPIClient(sg_api_key)
        response = sg.send(message)

        if 200 <= response.status_code < 300:
            return f"Success: Email sent to {to_email_addr} with {len(attachments)} attachment(s)."
        else:
            return f"Failure: SendGrid responded with status {response.status_code}."

    except Exception as e:
        return f"An error occurred: {e}"
    
# --- The updated tool definition ---
tool_send_email = StructuredTool.from_function(
    name="send_email_with_attachments",
    func=send_email_with_attachments,
    description=(
        "Use this tool to send an email with an XLSX timesheet and a DOCX invoice attached. "
        "It requires the filename for both documents ('xlsx_filename', 'docx_filename'). "
        "The files must be in the current working directory. "
        "The recipient's email ('to_email') is optional."
    )
)

tool_create_invoice_doc = StructuredTool.from_function(
    name="create_invoice_document",
    func=create_invoice_document,
    description=(
        "Use this to generate and save a formatted invoice as a Word (.docx) file. "
        "Input must be a dictionary with 'filename' and 'data' keys."
    )
)

tool_read_timesheet = Tool(
    name="read_invoice_data",
    func=read_invoice_data,
    description=(
        "Use this to read invoice data from an Excel timesheet file (XLSX format). "
        "It returns rows of date, status, and remarks. Input must be the filename (e.g., 'timesheet_july.xlsx')."
    )
)

tool_save_or_update_timesheet = StructuredTool.from_function(
    name="save_or_update_timesheet",
    func=save_or_update_timesheet,
    description=(
        "Use this to data to an Excel timesheet file (XLSX format). "
        "It needs the filename, date, status, and remarks as input. Filename usually be in the format: 'timesheet_<month>.xlsx')."
    )
)

tools = [tool_read_timesheet, tool_create_invoice_doc, tool_save_or_update_timesheet, tool_send_email]